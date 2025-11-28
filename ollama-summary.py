#the purpose of this automation tool is to summarize notes for students who wants to maximize their time in studying
#The tool will get the input as .pdf or .ppt and will return .docx with highlights and summary for better cognitive presentation
import sys
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
import subprocess
import requests
import time



#*************************************************************************
# EXTRACTION OF DOCUMENT
#*************************************************************************
#Function that converts pdf into text format
def pdfToText(pdf_path):
    openPDFFile = fitz.open(pdf_path)
    text = ""
    #store all the text inside the pdf
    for page in openPDFFile:
        text += page.get_text()
    return text

#Function that converts ppt presentation into text format
def pptToText(ppt_path : Path , pdfpath_out : Path = None):

    #Convert a PowerPoint (.pptx) file to PDF using LibreOffice (cross-platform).
    if pdfpath_out is None:
         pdfpath_out = ppt_path.parent

    pdf_out = pdfpath_out/"outputpdf"
    
    pdfOutput = pdf_out/(str(Path(ppt_path).stem) + ".pdf")
    print(f"Converting {Path(ppt_path).name} → {Path(pdfOutput).name} ...")

    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        str(ppt_path),
        "--outdir", str(pdf_out)
    ], check=True)

    print(f"Done: {pdfOutput}")
    text = pdfToText(pdfOutput)
    return text


#Function that will write the extracted information into .docx format
def createDocxFile(context, filename):
        createdDoc = Document()
        createdDoc.add_paragraph(context)
        createdDoc.save(filename)
        print(f"Saved: {filename}")

#**********************************************************************
# Summarize text for note taking and filtered information
#**********************************************************************
#Fuction that will read and extract all text from the document 
def readDocument(filename):
        readDocument = Document(filename)
        return "\n".join([paragraph.text for paragraph in readDocument.paragraphs])

#Function that will summarize context 
def summarizeContext(context):
        print("Initializing summarizer... it may take a while please wait...")
        
        # Load from cache only
        #tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn", local_files_only=True)
        #model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn", local_files_only=True)

        OLLAMA_URL = "http://192.168.1.18:11434/api/generate"   # change to your running ollama server for this situation my current server is running in a private network. 
        MODEL_NAME = "llama3:latest"  


        #Using llama3 model to perform the summarizer
        print("Using Ollama Llama model for summarization...")

        chunks = []
        # Limitation to prevent overflowing
        chunkSize = 900


        for i in range(0, len(context), chunkSize):
             chunks.append(context[i:i+ chunkSize])

        summaries_generated = []

        for i, chunk in enumerate(chunks):
             print(f" > Summarizing chunk {i+1}/{len(chunks)}")
             

             payload = {
                "model": MODEL_NAME,
                "prompt": f"""
                You are an expert summarization and note-taking assistant. 
                Your task is to read the following document text and extract the most critical information,
                concepts, and findings. Present your entire output as a single, concise list of bullet points.
                Return summary and concluding sentences without summarizing. Begin immediately with the first bullet point.

                Text:
                {chunk}
                """,
                "stream": False
            }
        
             try:
                response = requests.post(OLLAMA_URL, json=payload)
                data = response.json()

                #print(data) - debug

                # Ollama returns output under "response"
                summary_text = data.get("response", "").strip()

                summaries_generated.append(summary_text)
                time.sleep(0.2) # prevent too-fast calls to local server

             except Exception as e:
                print(f"Error summarizing chunk: {e}")
                summaries_generated.append("[ERROR IN CHUNK PROCESSING]")

        #return "\n\n".join(summaries)
        summaries = "\n\n".join(summaries_generated)
        return summaries



def main():
    if len(sys.argv) < 2:
        #Instruct user on the format of using the tool.
        print("\nUsage:")
        print("  python3 study-summary-automation.py extract\n")
        print("  python3 study-summary-automation.py summarize\n")
        sys.exit(1)

    instruction = sys.argv[1]
    #documentPath = sys.argv[2]
    base_path = Path.cwd()
    inputsDirectory = base_path /"selectedInputs"
    outputsDirectory = base_path/"output"

    #If the target folder does not exist
    if not inputsDirectory.exists():
        print(f"Folder not found: {inputsDirectory}")
        sys.exit(1)

    #choosing instruction 
    match instruction:
        case "extract":
            
            # Find all PDF and PPTX files
            files = [f for f in inputsDirectory.iterdir() if f.suffix.lower() in [".pdf", ".pptx"]]

            if not files:
                print("No PDF or PPTX files found in 'selectedInputs/'")
                sys.exit(0)

            #perform extraction to all the files that satisfies the requirement
            for file in files:
                print(f"Extracting from: {file.name}")

                #validates if file is either pdf or ppt
                if str(file).endswith(".pdf"):
                    context = pdfToText(str(file))
                    newFileName = file.name.replace(".pdf","_extracted.docx")

                #to be follow for the pptx presentation
                #********************************************************
                elif str(file).endswith(".pptx"):
                    context = pptToText(str(file), outputsDirectory)
                    newFileName = file.name.replace(".pptx","_extracted.docx")

                else:
                    print("Unsupported file format. Kindly use PDF to use the tool.")
                    sys.exit(1)

                print(f" Finished processing {file.name}\n")


                #paste all the extracted data into .docx format
               
                newFilePath = outputsDirectory/newFileName
                createDocxFile(context, outputsDirectory/f"{newFilePath}")
            
        
        case "summarize":
              # Find all DOCX files
            files = [f for f in inputsDirectory.iterdir() if f.suffix.lower() in [".docx"]]

            if not files:
                print("No DOCX files found in 'selectedInputs/'")
                sys.exit(0)

             #perform summarization to all the files that satisfies the requirement
            for file in files:
                print(f"Summarizing from: {file.name}")
                print(f"\nConverting {file.name} into study notes... Please wait...")

                context = readDocument(str(file))
                summarized = summarizeContext(context)

                newFileName = file.name.replace("_extracted.docx","_summary.docx")
                newFilePath = outputsDirectory/newFileName
                createDocxFile(summarized, f"{newFilePath}")

        case _: #default case - if neither of the two selection above
              print("Invalid command. Use either 'extract' or 'summarize'." )


if __name__ == "__main__":
    main()
